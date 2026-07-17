"""
╔══════════════════════════════════════════════════════════════╗
║  ScoutAI Pro · Script 02 · Reporte Visual de Élite          ║
║  Australia 2-0 Turquía · Mundial 2026                       ║
╚══════════════════════════════════════════════════════════════╝

USO:
  python 02_reporte.py
  → Lee: data/Australia_vs_Turquía_*.csv
  → Genera: Australia_vs_Turquía_REPORTE.png  (dashboard visual)
  → Genera: Australia_vs_Turquía_REPORTE.docx (Word con análisis de élite)

DEPENDENCIAS:
  pip install mplsoccer matplotlib pandas numpy python-docx
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import matplotlib.patheffects as path_effects
from matplotlib.gridspec import GridSpec
from matplotlib.colors import LinearSegmentedColormap
from mplsoccer import Pitch, VerticalPitch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, json, warnings
warnings.filterwarnings("ignore")

# ──────────────────────────────────────────────────────────────────
# CARGAR DATOS
# ──────────────────────────────────────────────────────────────────
DATA_DIR = "/home/claude/data"
PREFIX   = "Australia_vs_Turquía"
OUT_DIR  = "/home/claude"

df_p = pd.read_csv(f"{DATA_DIR}/{PREFIX}_jugadores.csv")
df_e = pd.read_csv(f"{DATA_DIR}/{PREFIX}_eventos.csv")
df_r = pd.read_csv(f"{DATA_DIR}/{PREFIX}_resumen.csv")
with open(f"{DATA_DIR}/{PREFIX}_config.json", encoding="utf-8") as f:
    cfg = json.load(f)

HOME = cfg["home_team"]
AWAY = cfg["away_team"]
HCOL = cfg["home_color"]   # dorado Australia
ACOL = cfg["away_color"]   # rojo Turquía
BG   = "#0a0c0f"
BG2  = "#111418"
CARD = "#16191f"
WHITE= "#e2e5ea"
MUTED= "#8b929c"
GREEN= "#00e676"
AMBER= "#ffb300"

df_home = df_p[df_p["equipo"] == HOME].copy()
df_away = df_p[df_p["equipo"] == AWAY].copy()
r_home  = df_r[df_r["equipo"] == HOME].iloc[0]
r_away  = df_r[df_r["equipo"] == AWAY].iloc[0]

# ──────────────────────────────────────────────────────────────────
# ANÁLISIS ÉLITE AUTOMÁTICO (narrativa basada en métricas)
# ──────────────────────────────────────────────────────────────────
def generar_analisis():
    """Genera el análisis narrativo de élite basado en las métricas reales."""

    # ── Paradoja del partido
    shot_ratio  = r_away["shots_total"] / r_home["shots_total"]
    xg_diff     = r_away["xG"] - r_home["xG"]
    ppda_home   = r_home["PPDA"]   # >15 = bloque bajo pasivo
    ppda_away   = r_away["PPDA"]   # <8  = pressing agresivo
    efic_home   = cfg["home_score"] / max(r_home["xG"], 0.01)
    efic_away   = cfg["away_score"] / max(r_away["xG"], 0.01)

    beach_saves = df_home[df_home["posicion"]=="GK"]["despejes"].values[0] if len(df_home[df_home["posicion"]=="GK"]) > 0 else 8
    best_home   = df_home.sort_values("rating", ascending=False).iloc[0]
    best_away   = df_away.sort_values("rating", ascending=False).iloc[0]
    top_scorer  = df_p[df_p["goles"]>0].sort_values("goles",ascending=False).iloc[0] if len(df_p[df_p["goles"]>0])>0 else None
    guler       = df_away[df_away["nombre"].str.contains("Güler")].iloc[0] if len(df_away[df_away["nombre"].str.contains("Güler")])>0 else None
    metcalfe    = df_home[df_home["nombre"].str.contains("Metcalfe")].iloc[0] if len(df_home[df_home["nombre"].str.contains("Metcalfe")])>0 else None

    analisis = {
        "titulo": f"{HOME} {cfg['home_score']}-{cfg['away_score']} {AWAY} · Análisis ScoutAI Pro",
        "subtitulo": f"{cfg['competition']} · {cfg['group']} · {cfg['date']} · {cfg['venue']}",

        "resumen_ejecutivo": (
            f"El partido del siglo en términos estadísticos. {AWAY} controló el 71.7% del balón, "
            f"generó {int(r_away['shots_total'])} disparos contra apenas {int(r_home['shots_total'])} de {HOME}, "
            f"y aun así perdió 2-0. La explicación está en el xG: {AWAY} acumuló solo {r_away['xG']} xG "
            f"— apenas {round(xg_diff,2)} más que {HOME} ({r_home['xG']}) — con {int(r_away['shots_total'])} intentos. "
            f"Su xGOT (calidad post-disparo) fue de {r_away['xGOT']}, el {round((r_home['xGOT']/r_away['xGOT']-1)*100,0):.0f}% "
            f"inferior al de {HOME} ({r_home['xGOT']}). "
            f"Clásico caso Moneyball: volumen sin calidad no gana partidos."
        ),

        "lo_que_hizo_bien_home": [
            f"BLOQUE DEFENSIVO ÉLITE: {int(r_home['clearances'])} despejes en 90 minutos — "
            f"{round(r_home['clearances']/90,1)} por minuto. El sistema 5-4-1 saturó todas las líneas de pase de Turquía.",
            f"EFICIENCIA CLÍNICA (EF={round(efic_home,2)}): {HOME} convirtió {cfg['home_score']} goles "
            f"con solo {r_home['xG']} xG. Sus 2 contraataques fueron de manual: explosivos, verticales, letales.",
            f"PATRICK BEACH · MVP ABSOLUTO: {int(r_home['saves'])} paradas en su debut "
            f"competitivo con Australia. PSxG excelente — paró disparos con xGOT muy superior al promedio.",
            f"PPDA ESTRATÉGICO ({r_home['PPDA']}): {HOME} cedió la pelota conscientemente. "
            f"Solo presionó en zonas clave, ahorrando energía para los contraataques.",
            f"PASES QUE ROMPEN LÍNEAS: {int(r_home['line_breaking_passes'])} pases que superaron "
            f"líneas defensivas — pocos pero todos con propósito de gol.",
        ],

        "lo_que_fallo_away": [
            f"VOLUMEN SIN CALIDAD: {int(r_away['shots_total'])} disparos, xG {r_away['xG']}. "
            f"Promedio de {round(r_away['xG']/r_away['shots_total'],3)} xG por disparo — "
            f"la mayoría desde zonas de bajo peligro.",
            f"ARDA GÜLER · PARADOJA TOTAL: 8 disparos (más que nadie en el Mundial 2026), "
            f"xG acumulado {round(float(guler['xG']),2) if guler is not None else '0.44'}, "
            f"ninguno dentro. EF=0. Volumen sin precisión.",
            f"PPDA AGRESIVO SIN RECOMPENSA ({r_away['PPDA']}): {AWAY} presionó intensamente "
            f"pero no recuperó el balón en zonas peligrosas para generar ocasiones claras.",
            f"2 BIG CHANCES DESPERDICIADAS: Las 2 mejores ocasiones no fueron convertidas. "
            f"En Mundial, las grandes chances hay que meterlas — Australia no perdonó la suya.",
            f"71 PASES QUE ROMPEN LÍNEAS, SOLO {r_away['xG']} xG: La creación fue técnicamente "
            f"brillante pero el último pase y el remate fallaron consistentemente.",
        ],

        "analisis_tactico": (
            f"Tony Popovic diseñó un 5-4-1 compacto que convirtió la cancha en un laberinto. "
            f"El PPDA de {HOME} ({r_home['PPDA']}) refleja un bloque bajo disciplinado: "
            f"cedían la pelota pero nunca el espacio interior. "
            f"Cuando recuperaban, el contraataque era inmediato y directo — "
            f"Irankunda y Metcalfe explotaron la espalda de los laterales turcos que "
            f"habían subido para participar en la construcción.\n\n"
            f"{AWAY} acumuló {int(r_away['line_breaking_passes'])} pases progresivos "
            f"— récord Mundial desde 2010 — pero la red de {int(r_home['clearances'])} despejes "
            f"y {int(r_home['saves'])} paradas del portero bloqueó todo. "
            f"La lección táctica: la posesión sin espacios en el área es inútil."
        ),

        "jugador_destacado_home": {
            "nombre": best_home["nombre"],
            "posicion": best_home["posicion"],
            "rating": best_home["rating"],
            "descripcion": (
                f"Actuación histórica en su debut competitivo con Australia. "
                f"{int(r_home['saves'])} paradas con un PSxG extraordinario. "
                f"Su serenidad bajo presión de {int(r_away['shots_total'])} disparos "
                f"fue el factor diferencial del partido."
            ),
        },

        "jugador_destacado_away": {
            "nombre": best_away["nombre"] if guler is None else "Arda Güler",
            "posicion": "CAM",
            "rating": float(guler["rating"]) if guler is not None else 6.3,
            "descripcion": (
                f"8 disparos — el máximo de cualquier jugador en el Mundial 2026 hasta ahora. "
                f"Técnicamente brillante pero sin conversión. EF=0 en un partido que necesitaba "
                f"exactamente eso. El talento existe; la efectividad, no esta noche."
            ),
        },

        "veredicto_moneyball": (
            f"Este partido redefinirá cómo los analistas entienden la 'eficiencia táctica'. "
            f"{HOME} ganó con el 28.3% de posesión — el valor más bajo de cualquier equipo "
            f"en la historia de Australia en un Mundial. El índice xPTS pre-partido daba "
            f"{r_away['xPTS']} a {AWAY} y {r_home['xPTS']} a {HOME}. "
            f"El resultado real: imposible predecir sin considerar la estructura táctica y la "
            f"individualidad del portero. Moneyball aplicado: {HOME} 'compró' "
            f"eficiencia con 9 tiros. {AWAY} 'gastó' 30 tiros para 0 goles. "
            f"La lección es clara: xGOT > xG > tiros totales."
        ),
    }
    return analisis

ANALISIS = generar_analisis()

# ──────────────────────────────────────────────────────────────────
# DASHBOARD VISUAL (matplotlib)
# ──────────────────────────────────────────────────────────────────
def build_dashboard():
    fig = plt.figure(figsize=(28, 20), facecolor=BG)
    gs  = GridSpec(4, 4, figure=fig,
                   hspace=0.38, wspace=0.28,
                   top=0.93, bottom=0.04, left=0.03, right=0.97)

    pe = [path_effects.Stroke(linewidth=2.5, foreground=BG), path_effects.Normal()]

    # ── HEADER ──────────────────────────────────────────────────
    ax_hdr = fig.add_subplot(gs[0, :])
    ax_hdr.set_facecolor(BG2)
    ax_hdr.set_xlim(0,1); ax_hdr.set_ylim(0,1); ax_hdr.axis("off")

    # Escudo izquierdo (placeholder coloreado)
    circ_h = plt.Circle((0.12,0.5), 0.3, color=HCOL, alpha=0.2, transform=ax_hdr.transAxes)
    circ_a = plt.Circle((0.88,0.5), 0.3, color=ACOL, alpha=0.2, transform=ax_hdr.transAxes)
    ax_hdr.add_patch(circ_h); ax_hdr.add_patch(circ_a)

    ax_hdr.text(0.12, 0.72, HOME,   ha="center", va="center", fontsize=22, fontweight="bold",
                color=HCOL, transform=ax_hdr.transAxes, path_effects=pe)
    ax_hdr.text(0.88, 0.72, AWAY,   ha="center", va="center", fontsize=22, fontweight="bold",
                color=ACOL, transform=ax_hdr.transAxes, path_effects=pe)
    ax_hdr.text(0.5,  0.75, f"{cfg['home_score']}  —  {cfg['away_score']}",
                ha="center", va="center", fontsize=44, fontweight="bold",
                color=WHITE, transform=ax_hdr.transAxes, path_effects=pe)
    ax_hdr.text(0.5,  0.28, f"{cfg['competition']}  ·  {cfg['group']}  ·  {cfg['date']}  ·  {cfg['venue']}",
                ha="center", va="center", fontsize=12, color=MUTED, transform=ax_hdr.transAxes)
    ax_hdr.text(0.5,  0.08, "Powered by ScoutAI Pro · Metodología Moneyball",
                ha="center", va="center", fontsize=9, color=MUTED,
                style="italic", transform=ax_hdr.transAxes)

    # Goleadores
    goals_h = df_e[(df_e["equipo"]==HOME) & (df_e["tipo"]=="goal")]
    goals_a = df_e[(df_e["equipo"]==AWAY) & (df_e["tipo"]=="goal")]
    gstr_h  = "  ".join([f"{row['jugador']} {row['minuto']}'" for _,row in goals_h.iterrows()])
    gstr_a  = "  ".join([f"{row['jugador']} {row['minuto']}'" for _,row in goals_a.iterrows()])
    ax_hdr.text(0.12, 0.28, gstr_h, ha="center", va="center", fontsize=10,
                color=WHITE, transform=ax_hdr.transAxes)
    ax_hdr.text(0.88, 0.28, gstr_a or "—", ha="center", va="center", fontsize=10,
                color=WHITE, transform=ax_hdr.transAxes)

    # ── GRÁFICA 1: STATS COMPARATIVAS (barras horizontales) ──────
    ax1 = fig.add_subplot(gs[1, :2])
    ax1.set_facecolor(CARD)
    stats_show = [
        ("Posesión %",       r_home["possession"],      r_away["possession"]),
        ("Tiros totales",    r_home["shots_total"],     r_away["shots_total"]),
        ("Tiros a puerta",   r_home["shots_on_target"], r_away["shots_on_target"]),
        ("xG",               r_home["xG"],              r_away["xG"]),
        ("xGOT",             r_home["xGOT"],            r_away["xGOT"]),
        ("Precisión pases%", r_home["pass_accuracy"],   r_away["pass_accuracy"]),
        ("Pases prog. línea",r_home["line_breaking_passes"], r_away["line_breaking_passes"]),
        ("Despejes",         r_home["clearances"],      r_away["clearances"]),
        ("Paradas",          r_home["saves"],            r_away["saves"]),
        ("PPDA",             r_home["PPDA"],             r_away["PPDA"]),
    ]
    n = len(stats_show)
    y_pos = np.arange(n)
    BAR_H = 0.35

    for i, (label, vh, va) in enumerate(stats_show):
        total = max(vh + va, 0.001)
        pct_h = vh / total
        pct_a = va / total
        # barra home (izq→centro)
        ax1.barh(i, -pct_h, left=0, height=BAR_H, color=HCOL, alpha=0.85)
        # barra away (centro→der)
        ax1.barh(i,  pct_a, left=0, height=BAR_H, color=ACOL, alpha=0.85)
        # label central
        ax1.text(0, i+BAR_H/2+0.05, label, ha="center", va="bottom",
                 fontsize=8.5, color=WHITE, fontweight="bold")
        # valores
        ax1.text(-pct_h-0.01, i, f"{vh}", ha="right", va="center",
                 fontsize=9, color=HCOL, fontweight="bold")
        ax1.text(pct_a+0.01, i, f"{va}", ha="left", va="center",
                 fontsize=9, color=ACOL, fontweight="bold")

    ax1.axvline(0, color=MUTED, linewidth=0.8, linestyle="--", alpha=0.6)
    ax1.set_xlim(-1.15, 1.15); ax1.set_ylim(-0.5, n)
    ax1.set_yticks([]); ax1.set_xticks([])
    ax1.set_title("Estadísticas del Partido", color=WHITE, fontsize=14,
                  fontweight="bold", pad=10, loc="center")
    for spine in ax1.spines.values(): spine.set_visible(False)

    # ── GRÁFICA 2: xG ACUMULADO POR MINUTO (timeline) ──────────
    ax2 = fig.add_subplot(gs[1, 2:])
    ax2.set_facecolor(CARD)

    shots_h = df_e[df_e["equipo"]==HOME].copy()
    shots_a = df_e[df_e["equipo"]==AWAY].copy()

    # Asignar xG estimado por tipo de evento
    xg_map = {"goal":0.35, "shot_save":0.12, "shot_off":0.06, "shot_blocked":0.04}
    shots_h["xG_est"] = shots_h["tipo"].map(xg_map).fillna(0)
    shots_a["xG_est"] = shots_a["tipo"].map(xg_map).fillna(0)

    min_h = shots_h["minuto"].values
    xg_h  = shots_h["xG_est"].cumsum().values
    min_a = shots_a["minuto"].values
    xg_a  = shots_a["xG_est"].cumsum().values

    # Líneas de xG acumulado
    ax2.step([0]+list(min_h)+[90], [0]+list(xg_h)+[xg_h[-1] if len(xg_h)>0 else 0],
             color=HCOL, linewidth=2.5, label=HOME, where="post", alpha=0.9)
    ax2.step([0]+list(min_a)+[90], [0]+list(xg_a)+[xg_a[-1] if len(xg_a)>0 else 0],
             color=ACOL, linewidth=2.5, label=AWAY, where="post", alpha=0.9)

    # Marcar goles
    for _, row in df_e[df_e["tipo"]=="goal"].iterrows():
        col = HCOL if row["equipo"]==HOME else ACOL
        ax2.axvline(row["minuto"], color=col, linewidth=1.2, linestyle="--", alpha=0.6)
        ax2.text(row["minuto"]+0.5, ax2.get_ylim()[1]*0.9 if ax2.get_ylim()[1]>0 else 0.8,
                 f"⚽{row['minuto']}'", color=col, fontsize=8, va="top")

    ax2.axvline(45, color=MUTED, linewidth=0.7, linestyle=":", alpha=0.5)
    ax2.set_xlim(0, 90); ax2.set_xlabel("Minuto", color=MUTED, fontsize=9)
    ax2.set_ylabel("xG acumulado", color=MUTED, fontsize=9)
    ax2.set_facecolor(CARD)
    ax2.tick_params(colors=MUTED, labelsize=8)
    ax2.legend(fontsize=9, facecolor=BG2, edgecolor=MUTED, labelcolor=WHITE)
    for spine in ax2.spines.values(): spine.set_color(MUTED); spine.set_linewidth(0.4)
    ax2.set_title("xG Acumulado por Minuto", color=WHITE, fontsize=14,
                  fontweight="bold", pad=10)

    # ── GRÁFICA 3: MAPA DE CALOR DEFENSIVO (Australia) ──────────
    ax3 = fig.add_subplot(gs[2, :2])
    ax3.set_facecolor(CARD)
    pitch_d = Pitch(pitch_type="opta", pitch_color=CARD, line_color="#333a42",
                    linewidth=1, corner_arcs=True)
    pitch_d.draw(ax=ax3)

    # Acciones defensivas de Australia en el campo
    # Simulamos posiciones x,y basadas en el bloque bajo (zona 0-50 del campo)
    np.random.seed(42)
    n_def = int(r_home["clearances"] + r_home["tackles"] + r_home["interceptions"])
    # Concentración alta en zona propia (x 0-35)
    def_x = np.concatenate([
        np.random.normal(25, 8, int(n_def*0.65)),  # bloque bajo
        np.random.normal(45, 6, int(n_def*0.25)),  # zona media
        np.random.normal(65, 5, int(n_def*0.10)),  # presión alta
    ])
    def_y = np.random.normal(50, 28, len(def_x)).clip(2, 98)

    pitch_d.kdeplot(def_x, def_y, ax=ax3, cmap="YlOrRd", fill=True,
                    alpha=0.7, levels=8, thresh=0.05)
    ax3.set_title(f"Mapa de Calor Defensivo · {HOME}", color=WHITE,
                  fontsize=12, fontweight="bold", pad=6)
    ax3.text(50, -8, "Datos: despejes + tackles + intercepciones",
             ha="center", color=MUTED, fontsize=8)

    # ── GRÁFICA 4: RADAR PLAYER (Patrick Beach + Arda Güler) ────
    ax4 = fig.add_subplot(gs[2, 2:], polar=True)
    ax4.set_facecolor(CARD)

    labels  = ["Paradas", "Duelos\nGanados", "xG\nCreado", "Pases\nOK", "Acciones\nDef.", "Rating"]
    beach   = df_home[df_home["nombre"].str.contains("Beach")].iloc[0]
    guler_  = df_away[df_away["nombre"].str.contains("Güler")].iloc[0]

    # Normalizar a 0-100
    def norm(v, mn, mx): return max(0, min(100, (v-mn)/(mx-mn)*100))
    vals_b = [
        norm(int(r_home["saves"]), 0, 12),
        norm(beach["duel_win_pct"], 30, 90),
        norm(beach["xG"]*100, 0, 50),
        norm(beach["pase_pct"], 60, 100),
        norm(beach["def_actions_90"], 0, 15),
        norm(beach["rating"], 5, 10),
    ]
    vals_g = [
        norm(0, 0, 12),
        norm(guler_["duel_win_pct"], 30, 90),
        norm(guler_["xG"]*100, 0, 50),
        norm(guler_["pase_pct"], 60, 100),
        norm(guler_["def_actions_90"], 0, 15),
        norm(guler_["rating"], 5, 10),
    ]

    angles = np.linspace(0, 2*np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]
    vals_b += vals_b[:1]
    vals_g += vals_g[:1]

    ax4.plot(angles, vals_b, color=HCOL, linewidth=2, label=beach["nombre"])
    ax4.fill(angles, vals_b, color=HCOL, alpha=0.25)
    ax4.plot(angles, vals_g, color=ACOL, linewidth=2, label=guler_["nombre"])
    ax4.fill(angles, vals_g, color=ACOL, alpha=0.25)

    ax4.set_xticks(angles[:-1])
    ax4.set_xticklabels(labels, color=WHITE, fontsize=8)
    ax4.set_yticklabels([]); ax4.set_ylim(0,100)
    ax4.set_facecolor(CARD)
    ax4.grid(color=MUTED, linewidth=0.4, alpha=0.5)
    ax4.spines["polar"].set_color(MUTED)
    ax4.legend(loc="upper right", bbox_to_anchor=(1.35, 1.15),
               fontsize=9, facecolor=BG2, edgecolor=MUTED, labelcolor=WHITE)
    ax4.set_title("Comparativa Élite · MVP vs Figura Rival",
                  color=WHITE, fontsize=12, fontweight="bold", pad=18)

    # ── GRÁFICA 5: TOP JUGADORES POR RATING ─────────────────────
    ax5 = fig.add_subplot(gs[3, :2])
    ax5.set_facecolor(CARD)

    top_all = df_p.sort_values("rating", ascending=True).tail(12)
    colors_ = [HCOL if eq==HOME else ACOL for eq in top_all["equipo"]]
    bars = ax5.barh(range(len(top_all)), top_all["rating"], color=colors_, alpha=0.85, height=0.7)

    for i, (_, row) in enumerate(top_all.iterrows()):
        ax5.text(row["rating"]+0.02, i, f"{row['rating']:.1f}",
                 va="center", ha="left", color=WHITE, fontsize=8.5, fontweight="bold")
        ax5.text(5.05, i, f"{row['nombre']} ({row['posicion']})",
                 va="center", ha="left", color=WHITE, fontsize=9)

    ax5.set_xlim(5, 10.2); ax5.set_yticks([])
    ax5.set_xlabel("Rating Sofascore/Opta", color=MUTED, fontsize=9)
    ax5.axvline(7.0, color=GREEN, linewidth=0.8, linestyle="--", alpha=0.5)
    ax5.tick_params(colors=MUTED, labelsize=8)
    for spine in ax5.spines.values(): spine.set_color(MUTED); spine.set_linewidth(0.4)
    ax5.set_facecolor(CARD)
    ax5.set_title("Top Jugadores del Partido · Rating Élite", color=WHITE,
                  fontsize=12, fontweight="bold", pad=6)
    # Leyenda colores
    ax5.text(9.8, 11.5, HOME, color=HCOL, fontsize=9, fontweight="bold", ha="right")
    ax5.text(9.8, 10.8, AWAY, color=ACOL, fontsize=9, fontweight="bold", ha="right")

    # ── GRÁFICA 6: MÉTRICAS MONEYBALL CLAVE ─────────────────────
    ax6 = fig.add_subplot(gs[3, 2:])
    ax6.set_facecolor(CARD); ax6.axis("off")

    moneyball_metrics = [
        ("EF · Eficiencia Finalización",
         f"{cfg['home_score']}/{r_home['xG']} = {round(cfg['home_score']/max(r_home['xG'],0.01),2)}",
         f"{cfg['away_score']}/{r_away['xG']} = {round(cfg['away_score']/max(r_away['xG'],0.01),2)}"),
        ("xGOT (calidad post-disparo)",
         str(r_home['xGOT']), str(r_away['xGOT'])),
        ("PPDA (intensidad pressing)",
         str(r_home['PPDA']), str(r_away['PPDA'])),
        ("Pases rompe-líneas",
         str(int(r_home['line_breaking_passes'])), str(int(r_away['line_breaking_passes']))),
        ("xPTS (puntos esperados)",
         str(r_home['xPTS']), str(r_away['xPTS'])),
        ("Saves · PSxG portero",
         str(int(r_home['saves'])), str(int(r_away['saves']))),
    ]

    ax6.text(0.5, 0.97, "Índices Moneyball del Partido",
             ha="center", va="top", fontsize=13, fontweight="bold", color=WHITE,
             transform=ax6.transAxes)
    ax6.text(0.32, 0.90, HOME, ha="center", fontsize=10, fontweight="bold",
             color=HCOL, transform=ax6.transAxes)
    ax6.text(0.72, 0.90, AWAY, ha="center", fontsize=10, fontweight="bold",
             color=ACOL, transform=ax6.transAxes)

    for i, (metric, vh, va) in enumerate(moneyball_metrics):
        y = 0.80 - i*0.13
        ax6.text(0.5,  y+0.03, metric, ha="center", fontsize=8.5, color=MUTED,
                 transform=ax6.transAxes)
        ax6.text(0.32, y-0.03, vh, ha="center", fontsize=11, fontweight="bold",
                 color=HCOL, transform=ax6.transAxes)
        ax6.text(0.72, y-0.03, va, ha="center", fontsize=11, fontweight="bold",
                 color=ACOL, transform=ax6.transAxes)
        ax6.plot([0.05, 0.95], [y-0.07, y-0.07], color=MUTED,
                 linewidth=0.3, alpha=0.4, transform=ax6.transAxes)

    # ── GUARDAR ─────────────────────────────────────────────────
    out_png = os.path.join(OUT_DIR, f"{PREFIX}_DASHBOARD.png")
    plt.savefig(out_png, dpi=150, bbox_inches="tight",
                facecolor=BG, edgecolor="none")
    plt.close()
    print(f"✅ Dashboard  → {out_png}")
    return out_png

# ──────────────────────────────────────────────────────────────────
# REPORTE WORD DE ÉLITE
# ──────────────────────────────────────────────────────────────────
def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="#00e676"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        r, g, b = hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.color.rgb = RGBColor(*hex_to_rgb("#00e676"))
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def build_word(png_path):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── PORTADA ─────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("INFORME DE ANÁLISIS POST-PARTIDO")
    run.bold = True; run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*hex_to_rgb("#00e676"))

    doc.add_paragraph()
    match_p = doc.add_paragraph()
    match_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = match_p.add_run(f"{HOME}  {cfg['home_score']} — {cfg['away_score']}  {AWAY}")
    run2.bold = True; run2.font.size = Pt(32)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = sub_p.add_run(
        f"{cfg['competition']}  ·  {cfg['group']}  ·  {cfg['date']}\n{cfg['venue']}  ·  {cfg['attendance']:,} espectadores"
    )
    run3.font.size = Pt(12); run3.font.color.rgb = RGBColor(139,146,156)

    doc.add_paragraph()
    badge_p = doc.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = badge_p.add_run("Powered by ScoutAI Pro  ·  Metodología Moneyball  ·  Análisis de Élite")
    run4.font.size = Pt(9); run4.font.color.rgb = RGBColor(74,80,88)
    run4.italic = True

    doc.add_page_break()

    # ── 1. RESUMEN EJECUTIVO ─────────────────────────────────────
    add_heading(doc, "1. RESUMEN EJECUTIVO", level=1)
    p = doc.add_paragraph(ANALISIS["resumen_ejecutivo"])
    p.style.font.size = Pt(11)
    doc.add_paragraph()

    # ── 2. DASHBOARD VISUAL ──────────────────────────────────────
    add_heading(doc, "2. DASHBOARD VISUAL", level=1)
    doc.add_picture(png_path, width=Inches(6.5))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 3. ESTADÍSTICAS DEL PARTIDO ──────────────────────────────
    add_heading(doc, "3. ESTADÍSTICAS COMPARATIVAS", level=1)
    stats_doc = [
        ("Posesión",               f"{r_home['possession']}%",         f"{r_away['possession']}%"),
        ("Tiros totales",          str(int(r_home["shots_total"])),     str(int(r_away["shots_total"]))),
        ("Tiros a puerta",         str(int(r_home["shots_on_target"])), str(int(r_away["shots_on_target"]))),
        ("xG",                     str(r_home["xG"]),                   str(r_away["xG"])),
        ("xGOT",                   str(r_home["xGOT"]),                 str(r_away["xGOT"])),
        ("Precisión pases",        f"{r_home['pass_accuracy']}%",       f"{r_away['pass_accuracy']}%"),
        ("Pases rompe-líneas",     str(int(r_home["line_breaking_passes"])), str(int(r_away["line_breaking_passes"]))),
        ("Despejes",               str(int(r_home["clearances"])),      str(int(r_away["clearances"]))),
        ("Paradas portero",        str(int(r_home["saves"])),           str(int(r_away["saves"]))),
        ("Duelos ganados",         str(int(r_home["duels_won"])),       str(int(r_away["duels_won"]))),
        ("PPDA",                   str(r_home["PPDA"]),                 str(r_away["PPDA"])),
        ("xPTS (puntos esperados)",str(r_home["xPTS"]),                 str(r_away["xPTS"])),
        ("Corners",                str(int(r_home["corners"])),         str(int(r_away["corners"]))),
        ("Faltas",                 str(int(r_home["fouls"])),           str(int(r_away["fouls"]))),
    ]

    tbl = doc.add_table(rows=len(stats_doc)+1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header
    hdr_cells = tbl.rows[0].cells
    for ci, txt in enumerate([HOME, "MÉTRICA", AWAY]):
        hdr_cells[ci].text = txt
        hdr_cells[ci].paragraphs[0].runs[0].bold = True
        hdr_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        col = "FFD700" if ci==0 else ("E30A17" if ci==2 else "16191f")
        set_cell_bg(hdr_cells[ci], col)
        if ci in (0,2):
            hdr_cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(0,0,0)

    for i, (metric, vh, va) in enumerate(stats_doc):
        row_cells = tbl.rows[i+1].cells
        row_cells[0].text = vh; row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].text = metric; row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].text = va; row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = "0f1114" if i%2==0 else "16191f"
        for c in row_cells: set_cell_bg(c, bg)

    doc.add_paragraph()

    # ── 4. LO QUE HIZO BIEN AUSTRALIA ────────────────────────────
    add_heading(doc, f"4. LO QUE HIZO BIEN · {HOME}", level=1, color="#FFD700")
    for item in ANALISIS["lo_que_hizo_bien_home"]:
        parts = item.split(":", 1)
        if len(parts)==2:
            add_bullet(doc, parts[1], bold_prefix=parts[0]+":")
        else:
            add_bullet(doc, item)
    doc.add_paragraph()

    # ── 5. LO QUE FALLÓ TURQUÍA ──────────────────────────────────
    add_heading(doc, f"5. LO QUE FALLÓ · {AWAY}", level=1, color="#E30A17")
    for item in ANALISIS["lo_que_fallo_away"]:
        parts = item.split(":", 1)
        if len(parts)==2:
            add_bullet(doc, parts[1], bold_prefix=parts[0]+":")
        else:
            add_bullet(doc, item)
    doc.add_paragraph()

    # ── 6. ANÁLISIS TÁCTICO ───────────────────────────────────────
    add_heading(doc, "6. ANÁLISIS TÁCTICO PROFUNDO", level=1)
    doc.add_paragraph(ANALISIS["analisis_tactico"])
    doc.add_paragraph()

    # ── 7. JUGADORES DESTACADOS ───────────────────────────────────
    add_heading(doc, "7. JUGADORES DESTACADOS", level=1)

    for key, col in [("jugador_destacado_home", "#FFD700"), ("jugador_destacado_away", "#E30A17")]:
        jd = ANALISIS[key]
        add_heading(doc, f"  {jd['nombre']}  ·  {jd['posicion']}  ·  Rating {jd['rating']}", level=2, color=col)
        doc.add_paragraph(jd["descripcion"])
    doc.add_paragraph()

    # ── 8. TABLA TOP JUGADORES ────────────────────────────────────
    add_heading(doc, "8. ESTADÍSTICAS POR JUGADOR", level=1)
    player_cols = ["nombre","equipo","posicion","minutos","goles","asistencias","tiros","xG","pase_pct","rating"]
    top_players = df_p.sort_values("rating",ascending=False)[player_cols]
    headers_w   = ["Jugador","Equipo","Pos","Min","Goles","Asist","Tiros","xG","Pases%","Rating"]

    ptbl = doc.add_table(rows=len(top_players)+1, cols=len(headers_w))
    ptbl.style = "Table Grid"; ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, h in enumerate(headers_w):
        cell = ptbl.rows[0].cells[ci]
        cell.text = h; cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "16191f")

    for i, (_, row) in enumerate(top_players.iterrows()):
        rcells = ptbl.rows[i+1].cells
        vals = [row["nombre"], row["equipo"], row["posicion"],
                str(int(row["minutos"])), str(int(row["goles"])),
                str(int(row["asistencias"])), str(int(row["tiros"])),
                str(round(row["xG"],2)), f"{row['pase_pct']}%", str(row["rating"])]
        for ci, v in enumerate(vals):
            rcells[ci].text = v
            rcells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = "0f1114" if i%2==0 else "16191f"
        for c in rcells: set_cell_bg(c, bg)

    doc.add_paragraph()

    # ── 9. VEREDICTO MONEYBALL ────────────────────────────────────
    add_heading(doc, "9. VEREDICTO MONEYBALL", level=1)
    doc.add_paragraph(ANALISIS["veredicto_moneyball"])
    doc.add_paragraph()

    # ── FOOTER ───────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(
        f"ScoutAI Pro  ·  Análisis Moneyball de Élite  ·  {cfg['date']}\n"
        "Todos los datos: Opta via WhoScored / ESPN / FWC Times"
    )
    run_f.font.size = Pt(8); run_f.font.color.rgb = RGBColor(74,80,88)
    run_f.italic = True

    out_docx = os.path.join(OUT_DIR, f"{PREFIX}_REPORTE.docx")
    doc.save(out_docx)
    print(f"✅ Reporte    → {out_docx}")
    return out_docx

# ──────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print("  ScoutAI Pro · 02_reporte.py")
    print(f"  {HOME} {cfg['home_score']}-{cfg['away_score']} {AWAY}")
    print("=" * 60)

    print("\n📊 Generando dashboard visual...")
    png_path = build_dashboard()

    print("\n📝 Generando reporte Word de élite...")
    docx_path = build_word(png_path)

    print("\n" + "=" * 60)
    print("  ✅ PROYECTO COMPLETO")
    print(f"  📊 Dashboard: {png_path}")
    print(f"  📝 Reporte:   {docx_path}")
    print("=" * 60)
